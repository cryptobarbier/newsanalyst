# -*- coding: utf-8 -*-
"""
Created on Sat Nov 23 19:34:15 2019

@author: PC
"""
import requests
import json
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel
from sklearn.cluster import AgglomerativeClustering
from summarizer import Summarizer
from docx import Document
import docx
from docx.shared import Pt

cat_to_drop=['dmoz/Business/Investing/Day Trading','dmoz/Shopping/Gifts','news/Arts and Entertainment','dmoz/Home/Personal_Finance/Tax_Preparation','dmoz/Recreation/Travel/Transportation','dmoz/Shopping/Holidays','dmoz/Recreation/Travel/Transportation','news/Sports','dmoz/Computers/Hardware/Peripherals','dmoz/Business/Employment/Job Search','dmoz/Business/Investing/Guides','dmoz/Society/Work/Work and Family']


#Extract the site rank of json result
def Rank(sce):
  try:
    return sce['ranking']['alexaGlobalRank']
  except:
    return 1500000

# Extract Number of Facebook shares
def Shares(sh):
  try:
    return sh['facebook']
  except:
    return 0

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

class NewsArticle():
    def __init__(self,req,cutoff_ranking=30000,weight_cutoff=20,max_len=600):
            #Run the request and put the results into a json file
        res=requests.get(req)
        j=json.loads(res.text[14:-1],encoding='ascii')['articles']['results']
        df = pd.DataFrame.from_dict(j, orient='columns')
        #Remove irrelevant articles
        df=df[df['wgt']>weight_cutoff]
        #Remove Duplicates
        df=df[df['isDuplicate']==False]
        # Remove duplicate bodys i.e articles published in several websites
        df.drop_duplicates(subset ='title', 
                             keep = 'first', inplace = True) 
        # Filter the articles on source rating
        df['Source Rank']=df['source'].apply(Rank)
        df['Shares']=df['shares'].apply(Shares)
        # replace the /n
        df['body']=df['body'].str.replace('\n',' ')
        df['body']=df['body'].str.replace("\'s","'s")
        df.sort_values(by='Source Rank')
        df=df[df['Source Rank']<cutoff_ranking]
        self.results=df
        self.max_len=max_len
    
    def CreateDist(self):
        
        # Build the text database
        corp=[self.results.loc[x]['body'] for x in self.results.index ]
        tfidf = TfidfVectorizer().fit_transform(corp)
        cosine_similarities = linear_kernel(tfidf, tfidf)
        matrix=pd.DataFrame(cosine_similarities,index=self.results.index,columns=self.results.index)
        self.dist=1-matrix
        
    def Cluster(self,thresh=0.4):
        clustering = AgglomerativeClustering(affinity='precomputed',distance_threshold=thresh,linkage='complete',n_clusters=None).fit(self.dist)
        self.clust=pd.DataFrame(clustering.labels_,index=self.results.index).sort_values(by=0)
        
    def Cluster2(self,thresh=0.42):
        cluster_new=pd.DataFrame(index=self.dist.index,columns=[0])
        d1=self.dist[self.dist<thresh]
        cluster_nb=0
        while len(d1)>0:
            cl_indices=d1.iloc[0].dropna().index
            for i in cl_indices:
              cluster_new.at[cl_indices,0]=cluster_nb
            cluster_nb=cluster_nb+1
            d1=d1.drop(cl_indices)
            d1=d1.drop(cl_indices,axis=1)
        self.clust=cluster_new
    
    def Summary(self):
        summ=dict.fromkeys(self.clust[0].unique())
        # Will store the name of Cluster based on article name
        subtitles=[]
        # Will store the url of main article summarizing the cluster
        urls=[]
        model = Summarizer()
        # Create the summary by aggregating all the bodys of the articles
        for c in summ.keys():
            df2=self.clust[self.clust[0]==c]
            subtitles=subtitles+[self.results.loc[df2.index[0]]['title']]
            urls=urls+[self.results.loc[df2.index[0]]['url']]
            full_text=' '.join(self.results.loc[df2.index]['body'].unique())
            full_text=full_text.replace('\n',' ')
            length=len(full_text)
            if length<500:
              target=200
              ratiol=min(1,target/length)
            else:
              target=min(self.max_len+150*(len(df2)),length/len(df2)*(0.15+0.05*len(df2)))
              ratiol=min(1,target/length)
            summ[c]=model(full_text,min_length=60,ratio=ratiol)
        self.summary=summ
        self.subtitles=subtitles
        self.urls=urls



    def TextOutput(self,title,docname):
          document = Document()
          style = document.styles['Normal']
          font = style.font
          font.name = 'Calibri'
          font.size = Pt(10)
          # Write the title passed at the function
          document.add_heading(title, 0)
          i=0
          for element in self.summary:
            document.add_heading(self.subtitles[i], level=1)
            p = document.add_paragraph()
            #add a hyperlink with the normal formatting (blue underline)
            hyperlink = add_hyperlink(p, self.urls[i], self.urls[i], None, True)
            p = document.add_paragraph(self.summary[element])
            i=i+1
          document.save(docname+'.docx')
          self.document=document
    
    def CleanFinance(self):
        self.results['Cat']=self.results['categories'].apply(FilterCat)
        self.results=self.results[self.results['Cat']==1]

        # Save Ft articles in full 
    def FTSave(self,title,docname):
        ft_articles=self.results[self.results['url'].str[:18]=='https://www.ft.com']
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        # Write the title passed at the function
        document.add_heading(title, 0)
        i=0
        for element in ft_articles.index:
            document.add_heading(ft_articles.at[element,'title'], level=1)
            p = document.add_paragraph(ft_articles.at[element,'body'])
            i=i+1
        document.save(docname+'.docx')
        self.document=document


          
# Returns 1 if articles does not contain any category to drop
def FilterCat(categories):
    df_filter=pd.DataFrame.from_dict(categories)
    labs=df_filter['label'].unique()
    overlap=set(labs).intersection(cat_to_drop)
    if len(overlap)==0:
        return 1
    else:
        return 0


        
    
        

    
            